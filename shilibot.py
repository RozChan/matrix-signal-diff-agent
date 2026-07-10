"""
飞书知识库问答机器人
监听用户消息，从飞书知识库中搜索相关文档并回复摘要。

使用前请确保：
1. 应用已开启「机器人」能力
2. 已订阅 im.message.receive_v1 事件
3. 已申请 im:message.p2p_msg:readonly / im:message:send_as_bot / search:docs:read 权限
4. 已通过 lark-cli auth login 完成授权
"""

import json
import subprocess
import sys
import os
import re
import threading
import logging
import tempfile

# ─── 配置 ───────────────────────────────────────────────────────────────────
LARK_CLI = r"C:\Users\ch225078\AppData\Roaming\npm\node_modules\@larksuite\cli\bin\lark-cli.exe"

# 知识库空间 ID 列表（留空则搜索全部可访问文档）
# 通过 lark-cli wiki +space-list --as user --format json 查看可用空间
WIKI_SPACE_IDS: list[str] = []

# 固定知识文档列表（优先从这些文档中获取答案）
# 格式: {"token": "文档token", "title": "文档标题"}
#   - 普通飞书文档: 填 token
#   - 上传的文件(is_file): 优先用 local_file 读取本地副本，无本地副本时用 token 下载
PINNED_DOCS: list[dict] = [
    # {"token": "ZrrcdHNPcoLyR8xISs6cYR2Enqh", "title": "智慧能源之JIRA不懂就来问"},  # 暂时停用
    # {"title": "智慧能源Jira和Confluence常见问题", "local_file": r"f:\AI\智慧能源Jira和Confluence常见问题.docx"},  # 本地文件，已取消引用
    {"token": "TgfIwHZuoiC5xGkr0gzcbG5nnre", "title": "智慧能源Jira和Confluence常见问题", "url": "https://my-ichery.feishu.cn/wiki/TgfIwHZuoiC5xGkr0gzcbG5nnre"},
]

# 每次回复最多引用几篇文档
MAX_DOCS = 3

# 摘要最大字符数（每篇）
MAX_SUMMARY_CHARS = 30000

# 人工专家 open_id（机器人无法回答时转发给此人）
HUMAN_EXPERT_OPEN_ID = "ou_8491434975006a2f5cfc3ac2c29273e6"  # 汪鹏

# 知识库文档配置（专家回复纯文字问题后自动追加问答到此文档）
KB_DOC_TOKEN = "TgfIwHZuoiC5xGkr0gzcbG5nnre"

# 每条回复底部附加的转人工提示
TRANSFER_FOOTER = "\n\n---\n若回复未能解答您的疑问，请直接回复「**转人工**」，我们将为您转交专家处理。"

# 待人工处理的请求：{发送给专家的消息ID: {"questioner_id": 原提问者open_id, "question": 原始问题}}
# 以 sent_msg_id 为 key，支持多个用户同时待处理，专家通过 reply_to 精确匹配
_pending_expert_requests: dict[str, dict] = {}
_pending_lock = threading.Lock()

# 记录每个用户最近一次提问的问题：{sender_id: question}
_last_user_question: dict[str, str] = {}

# 记录每个用户最近一次发送的图片：{sender_id: {"image_key": key, "message_id": msg_id}}
_last_user_image: dict[str, dict] = {}

# 知识库文档操作锁（防止并发追加冲突）
_kb_lock = threading.Lock()

# 消息去重集合，防止同一事件被重复处理
_processed_messages: set[str] = set()
_processed_lock = threading.Lock()
_MAX_PROCESSED = 1000  # 最多保留多少条历史记录

# ─── 日志 ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    stream=sys.stderr,
)
log = logging.getLogger("kb-bot")


# ─── 工具函数 ──────────────────────────────────────────────────────────────
def run_cli(*args: str, timeout: int = 30) -> str | None:
    """运行 lark-cli 命令并返回 stdout，失败返回 None。"""
    cmd = [LARK_CLI, *args]
    log.debug("exec: %s", " ".join(cmd))
    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=timeout, encoding="utf-8"
        )
        if result.returncode != 0:
            log.warning("lark-cli failed (code=%d): %s", result.returncode, result.stderr.strip())
            return None
        return result.stdout.strip()
    except subprocess.TimeoutExpired:
        log.warning("lark-cli timeout after %ds", timeout)
        return None
    except Exception as e:
        log.error("lark-cli error: %s", e)
        return None


def search_docs(query: str) -> list[dict]:
    """在飞书知识库中搜索文档，返回匹配列表。"""
    search_args = ["drive", "+search", "--query", query, "--format", "json", "--as", "bot"]
    if WIKI_SPACE_IDS:
        search_args += ["--space-ids", ",".join(WIKI_SPACE_IDS)]

    output = run_cli(*search_args)
    if not output:
        return []

    try:
        data = json.loads(output)
        items = data.get("items") or data.get("data", {}).get("items") or data.get("files") or []
        # 只取 docx 类型
        docs = []
        for item in items:
            doc_type = item.get("type") or item.get("doc_type") or ""
            token = item.get("token") or item.get("file_token") or ""
            title = item.get("title") or item.get("name") or ""
            url = item.get("url") or ""
            if token and doc_type in ("docx", "doc", "wiki", ""):
                docs.append({"token": token, "title": title, "url": url, "type": doc_type})
            if len(docs) >= MAX_DOCS:
                break
        return docs
    except json.JSONDecodeError:
        log.warning("search result not valid JSON")
        return []


def fetch_doc_content(token: str) -> str:
    """获取文档正文内容（im-markdown 格式，适合直接发送）。"""
    output = run_cli(
        "docs", "+fetch",
        "--doc", token,
        "--doc-format", "im-markdown",
        "--as", "bot",
        timeout=20,
    )
    if output:
        try:
            data = json.loads(output)
            content = data.get("data", {}).get("document", {}).get("content", "")
            if content:
                return content
        except json.JSONDecodeError:
            pass
    # fallback: 尝试普通格式
    output = run_cli(
        "docs", "+fetch",
        "--doc", token,
        "--as", "bot",
        timeout=20,
    )
    if output:
        try:
            data = json.loads(output)
            content = data.get("data", {}).get("document", {}).get("content", "")
            if content:
                return content
        except json.JSONDecodeError:
            pass
    return ""


# 文件内容缓存（避免每次请求都重复读取）
_file_cache: dict[str, str] = {}


def read_local_docx(file_path: str) -> str:
    """读取本地 .docx 文件并提取文本内容（带缓存）。"""
    if file_path in _file_cache:
        return _file_cache[file_path]
    if not os.path.exists(file_path):
        log.warning("本地文件不存在: %s", file_path)
        return ""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style and para.style.name and "Heading" in para.style.name:
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        level = int(level)
                    except ValueError:
                        level = 2
                    text_parts.append(f"{'#' * level} {para.text}")
                else:
                    text_parts.append(para.text)
        content = "\n\n".join(text_parts)
        _file_cache[file_path] = content
        log.info("本地文件 [%s] 读取成功 (%d 字符)", os.path.basename(file_path), len(content))
        return content
    except Exception as e:
        log.error("读取本地 .docx 文件失败 [%s]: %s", file_path, e)
        return ""


def fetch_file_content(file_token: str) -> str:
    """下载飞书文件类型的 .docx 并提取文本内容。"""
    if file_token in _file_cache:
        return _file_cache[file_token]

    # 下载文件
    tmp_dir = tempfile.mkdtemp()
    output_path = os.path.join(tmp_dir, "doc.docx")
    # lark-cli 要求 --output 为相对路径，先切换到 tmp_dir
    old_cwd = os.getcwd()
    try:
        os.chdir(tmp_dir)
        output = run_cli(
            "drive", "+download",
            "--file-token", file_token,
            "--output", "doc.docx",
            "--as", "bot",
            timeout=30,
        )
    finally:
        os.chdir(old_cwd)
    if not output:
        log.warning("文件下载失败: %s", file_token)
        return ""

    # 检查文件是否存在
    local_path = output_path
    if not os.path.exists(local_path):
        # 尝试在 tmp_dir 中查找 .docx 文件
        for f in os.listdir(tmp_dir):
            if f.endswith(".docx"):
                local_path = os.path.join(tmp_dir, f)
                break
        else:
            log.warning("未找到下载的文件: %s", file_token)
            return ""

    # 解析 .docx 文件
    try:
        from docx import Document
        doc = Document(local_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 保留标题层级
                if para.style and para.style.name and "Heading" in para.style.name:
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        level = int(level)
                    except ValueError:
                        level = 2
                    text_parts.append(f"{'#' * level} {para.text}")
                else:
                    text_parts.append(para.text)
        content = "\n\n".join(text_parts)
        _file_cache[file_token] = content
        log.info("文件 [%s] 解析成功 (%d 字符)", file_token, len(content))
        return content
    except Exception as e:
        log.error("解析 .docx 文件失败: %s", e)
        return ""
    finally:
        # 清理临时文件
        try:
            for f in os.listdir(tmp_dir):
                os.remove(os.path.join(tmp_dir, f))
            os.rmdir(tmp_dir)
        except Exception:
            pass


def truncate(text: str, max_len: int = MAX_SUMMARY_CHARS) -> str:
    """截断文本到指定长度。"""
    text = text.strip()
    if len(text) <= max_len:
        return text
    return text[:max_len].rsplit("\n", 1)[0] + "\n...(内容已截断)"


# 需要过滤掉的无意义词（停用词）
STOP_WORDS = set("的了是在怎么办怎么什么为什么如何哪些哪里有没有可以不可以对和与或及但如果不就都也会能被把给从到为让我你他她它们着过来去做看想说吗呢吧啊请问一下这个那个怎样进行通过之后然后以后之前以前时候现在已经还是或者比较关于知道告诉帮该要得很我应该怎么开始使始使")

# 同义词映射：key 是标准词，value 是它的所有同义词
# 当用户提问包含某个同义词时，自动扩展为标准词一起参与匹配
SYNONYMS = {
    "网址": ["地址", "链接", "入口"],
    "账号": ["帐号"],
    "登录": ["登陆"],
}
# 构建反向映射：同义词 → 标准词
_SYNONYM_MAP = {}
for _std, _syns in SYNONYMS.items():
    for _syn in _syns:
        _SYNONYM_MAP[_syn] = _std


def expand_synonyms(keywords: list[str]) -> list[str]:
    """对关键词列表做同义词扩展。
    
    如果关键词中包含同义词，将对应的标准词也加入列表。
    例如：用户问"jira地址" → 扩展出"网址"。
    """
    expanded = list(keywords)
    seen = set(keywords)
    for kw in keywords:
        # 检查关键词本身是否是同义词
        if kw in _SYNONYM_MAP:
            std = _SYNONYM_MAP[kw]
            if std not in seen:
                expanded.append(std)
                seen.add(std)
        # 检查关键词是否包含同义词子串
        for syn, std in _SYNONYM_MAP.items():
            if syn in kw and std not in seen:
                expanded.append(std)
                seen.add(std)
    return expanded


def _all_chars_nonstop(text: str) -> bool:
    """检查文本中每个字符都不在停用词表中。"""
    return all(c not in STOP_WORDS for c in text)


def extract_keywords(text: str) -> list[str]:
    """从文本中提取关键词。
    
    1. 按非文字字符拆分
    2. 对中文做 bigram/trigram 拆分，仅保留每个字都非停用词的组合
    3. 短查询（≤8字）额外保留原始完整文本
    4. 去停用词
    5. 同义词扩展
    """
    import re
    # 先按 ASCII 分隔符拆分
    tokens = re.split(r'[\s,.\?!;:\-+=*/\\()\[\]{}]+', text.lower())
    keywords = []
    for token in tokens:
        token = token.strip()
        if len(token) < 2 or token in STOP_WORDS:
            continue
        # 检测是否含中文字符
        if any('\u4e00' <= c <= '\u9fff' for c in token):
            # 短查询保留完整原文
            if len(token) <= 8 and token not in STOP_WORDS:
                keywords.append(token)
            # 中文 bigram 拆分（仅保留两个字都不是停用词的组合）
            for i in range(len(token) - 1):
                bigram = token[i:i+2]
                if _all_chars_nonstop(bigram):
                    keywords.append(bigram)
            # trigram（仅保留三个字都不是停用词的组合）
            for i in range(len(token) - 2):
                trigram = token[i:i+3]
                if _all_chars_nonstop(trigram):
                    keywords.append(trigram)
        else:
            keywords.append(token)
    # 去重保序
    seen = set()
    result = []
    for kw in keywords:
        if kw not in seen:
            seen.add(kw)
            result.append(kw)
    # 同义词扩展
    result = expand_synonyms(result)
    return result


def extract_relevant_paragraphs(doc_content: str, question: str, max_paragraphs: int = 1, min_score: float = 0.35) -> str:
    """从文档中提取与问题最相关的段落。
    
    1. 按标题或双换行拆分为段落
    2. 将段落按问题编号（如“3、”“4、”）分组为完整问答块
    3. 计算每个问答块与问题的关键词匹配度
    4. 只返回得分 >= min_score 的块（强相关）
    5. 无强相关块时返回空字符串
    """
    # 拆分段落：按 Markdown 标题行或连续空行分割
    import re
    raw_blocks = re.split(r'\n(?=#{1,6}\s)|\n\s*\n', doc_content)
    paragraphs = []
    current_heading = ""
    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue
        heading_match = re.match(r'^(#{1,6})\s+(.+)', block)
        if heading_match:
            current_heading = heading_match.group(2).strip()
            rest = block[heading_match.end():].strip()
            if rest:
                paragraphs.append({"heading": current_heading, "text": rest})
        else:
            paragraphs.append({"heading": current_heading, "text": block})

    if not paragraphs:
        log.info("段落列表为空")
        return ""

    log.info("共解析到 %d 个段落", len(paragraphs))

    # 将段落分组为问答块（以“数字+、”或“数字+.”开头的段落为新块起点）
    blocks = []  # 每个 block = {"heading": str, "texts": [str]}
    for p in paragraphs:
        if re.match(r'^[A-Za-z]?\d+[、.]', p['text']):
            # 新问题编号，开新块
            blocks.append({"heading": p['heading'], "texts": [p['text']]})
        elif blocks:
            blocks[-1]["texts"].append(p['text'])
        # 如果没有块且不是问题编号（如元数据），跳过

    if not blocks:
        log.info("问答块列表为空，段落前10个: %s", [p['text'][:30] for p in paragraphs[:10]])
        return ""

    # 调试：打印所有问答块
    for i, b in enumerate(blocks):
        log.info("块[%d] 共%d段: %s", i, len(b['texts']), [t[:30] for t in b['texts']])

    # 提取问题关键词
    q_keywords = extract_keywords(question)
    log.info("问题关键词: %s", q_keywords)

    # 提取查询中有意义的字符（非停用词字符），用于字符级回退匹配
    q_meaningful_chars = set()
    for token in re.split(r'[\s,.\?!;:\-+=*/\\()\[\]{}]+', question.lower()):
        for c in token:
            if c not in STOP_WORDS and ('\u4e00' <= c <= '\u9fff' or c.isalpha()):
                q_meaningful_chars.add(c)
    char_overlap_threshold = 0.7

    if not q_keywords and not q_meaningful_chars:
        return ""

    # 为每个问答块打分（按关键词覆盖率 + 字符级回退）
    total_keywords = len(q_keywords)
    scored = []
    for idx, b in enumerate(blocks):
        combined = " ".join(b['texts']).lower()
        # n-gram 关键词匹配
        matched = sum(1 for kw in q_keywords if kw in combined) if total_keywords > 0 else 0
        score = matched / total_keywords if total_keywords > 0 else 0

        # 字符级回退：如果有意义的字符大部分出现在目标中，视为匹配
        if score < min_score and q_meaningful_chars:
            char_overlap = sum(1 for c in q_meaningful_chars if c in combined) / len(q_meaningful_chars)
            if char_overlap >= char_overlap_threshold:
                score = char_overlap
                log.info("块[%d] 字符回退 score=%.2f (%d/%d chars) text='%s'",
                         idx, score, int(char_overlap * len(q_meaningful_chars)),
                         len(q_meaningful_chars), b['texts'][0][:40])

        if score > 0:
            if score >= min_score or matched > 0:
                scored.append((score, idx, b))
                log.info("块[%d] score=%.2f (%d/%d) text='%s'", idx, score, matched, total_keywords, b['texts'][0][:40])

    if not scored:
        log.info("无匹配块")
        return ""

    scored.sort(key=lambda x: x[0], reverse=True)
    top = [(s, idx, b) for s, idx, b in scored if s >= min_score][:max_paragraphs]

    if not top:
        log.info("无强相关块（最高分=%.2f，阈值=%.2f）", scored[0][0], min_score)
        return ""

    log.info("匹配到 %d 个强相关块: %s", len(top), [(round(s, 2), b['texts'][0][:20]) for s, _, b in top])

    # 返回完整问答块（块内段落用 \n\n 连接）
    result_parts = []
    for score, idx, b in top:
        result_parts.append("\n\n".join(b['texts']))

    return "\n\n".join(result_parts)


def reply_message(message_id: str, text: str) -> bool:
    """回复一条消息。"""
    output = run_cli(
        "im", "+messages-reply",
        "--message-id", message_id,
        "--markdown", text,
        "--as", "bot",
    )
    return output is not None


def reply_text(message_id: str, text: str) -> bool:
    """回复一条纯文本消息。"""
    output = run_cli(
        "im", "+messages-reply",
        "--message-id", message_id,
        "--text", text,
        "--as", "bot",
    )
    return output is not None


def send_text_to_user(user_id: str, text: str) -> str | None:
    """主动给指定用户发一条私聊文本消息，返回 message_id，失败返回 None。"""
    output = run_cli(
        "im", "+messages-send",
        "--user-id", user_id,
        "--text", text,
        "--as", "bot",
        "--format", "json",
    )
    if output:
        try:
            data = json.loads(output)
            return data.get("message_id", "") or data.get("data", {}).get("message_id", "")
        except json.JSONDecodeError:
            pass
    return None


def send_markdown_to_user(user_id: str, markdown: str) -> str | None:
    """主动给指定用户发一条 markdown 消息（支持嵌入图片），返回 message_id。"""
    output = run_cli(
        "im", "+messages-send",
        "--user-id", user_id,
        "--markdown", markdown,
        "--as", "bot",
        "--format", "json",
    )
    if output:
        try:
            data = json.loads(output)
            return data.get("message_id", "") or data.get("data", {}).get("message_id", "")
        except json.JSONDecodeError:
            pass
    return None


def content_has_image(content: str) -> bool:
    """检查消息内容是否包含图片引用。
    支持两种格式：
    - ![Image](img_xxx-xxx)  — 嵌入图片的富文本
    - [Image: img_xxx-xxx]   — 纯图片消息
    """
    return bool(
        re.search(r'!\[.*?\]\(img_[a-zA-Z0-9_-]+\)', content)
        or re.search(r'\[Image:\s*img_[a-zA-Z0-9_-]+\]', content)
    )


def strip_image_markdown(content: str) -> str:
    """去除消息内容中的图片引用标记，返回纯文本。
    支持两种格式：
    - ![Image](img_xxx-xxx)  — 嵌入图片的富文本
    - [Image: img_xxx-xxx]   — 纯图片消息
    """
    # 去掉 ![xxx](img_xxx) 格式
    cleaned = re.sub(r'!\[.*?\]\(img_[a-zA-Z0-9_-]+\)', '', content)
    # 去掉 [Image: img_xxx] 格式
    cleaned = re.sub(r'\[Image:\s*img_[a-zA-Z0-9_-]+\]', '', cleaned)
    # 去掉多余的空白行
    cleaned = re.sub(r'\n{2,}', '\n', cleaned).strip()
    return cleaned


def get_user_name(open_id: str) -> str:
    """通过 open_id 查询用户姓名，失败时返回 open_id 本身。"""
    # 优先用 +search-user（user身份），能获取 localized_name
    output = run_cli(
        "contact", "+search-user",
        "--user-ids", open_id,
        "--as", "user",
        "--format", "json",
    )
    if output:
        try:
            data = json.loads(output)
            users = data.get("data", {}).get("users", [])
            if users:
                name = users[0].get("localized_name", "") or users[0].get("name", "")
                if name:
                    return name
        except json.JSONDecodeError:
            pass
    # 回退：用 +get-user（bot身份）
    output = run_cli(
        "contact", "+get-user",
        "--user-id", open_id,
        "--as", "bot",
        "--format", "json",
    )
    if output:
        try:
            data = json.loads(output)
            name = data.get("data", {}).get("user", {}).get("name", "")
            if name:
                return name
        except json.JSONDecodeError:
            pass
    return open_id


def check_message_reply(message_id: str) -> str | None:
    """检查消息是否引用了另一条消息，返回被引用的 message_id，无引用返回 None。"""
    output = run_cli(
        "im", "+messages-mget",
        "--message-ids", message_id,
        "--as", "bot",
        "--format", "json",
    )
    if output:
        try:
            data = json.loads(output)
            messages = data.get("data", {}).get("messages", [])
            if messages:
                return messages[0].get("reply_to") or None
        except json.JSONDecodeError:
            pass
    return None


def _fetch_kb_xml() -> str:
    """获取知识库文档 XML 内容（含 block ID），失败返回空字符串。"""
    output = run_cli(
        "docs", "+fetch",
        "--doc", KB_DOC_TOKEN,
        "--detail", "with-ids",
        "--as", "bot",
        "--format", "json",
        timeout=20,
    )
    if output:
        try:
            data = json.loads(output)
            return data.get("data", {}).get("document", {}).get("content", "")
        except (json.JSONDecodeError, ValueError):
            pass
    return ""


def _find_insert_point(xml_content: str) -> tuple[int, str]:
    """从 XML 内容中找到最后一个完整条目的答案段落 block ID。
    返回 (最大编号, answer_block_id)，未找到返回 (0, "")。
    
    逻辑：找到最后一个编号段落（问题），然后找到紧跟它的下一个非空段落（答案），
    返回答案段落的 block_id 作为插入点，这样新条目会插入在旧条目答案之后。
    """
    # 匹配所有编号段落：<p id="xxx">可选空白+数字、
    q_pattern = r'<p\s+id="([^"]+)"[^>]*>\s*(\d+)\u3001'
    q_matches = list(re.finditer(q_pattern, xml_content))
    if not q_matches:
        return 0, ""
    
    # 找最大编号
    best_num = 0
    best_q_id = ""
    best_match_end = 0
    for m in q_matches:
        num = int(m.group(2))
        if num > best_num:
            best_num = num
            best_q_id = m.group(1)
            best_match_end = m.end()
    
    # 在问题段落之后找到下一个段落（答案）
    # 模式：<p id="xxx">非空内容</p>
    answer_pattern = r'<p\s+id="([^"]+)"[^>]*>([^<]+)</p>'
    m_answer = re.search(answer_pattern, xml_content[best_match_end:])
    
    if m_answer and m_answer.group(2).strip():
        # 找到答案段落，返回其 block_id
        answer_id = m_answer.group(1)
        log.debug("找到 #%d 的答案段落 block=%s", best_num, answer_id)
        return best_num, answer_id
    else:
        # 没有找到答案段落（可能是纯问题无答案），返回问题段落 block_id
        log.debug("未找到 #%d 的答案段落，使用问题段落 block=%s", best_num, best_q_id)
        return best_num, best_q_id


def _insert_qa_with_retry(last_block_id: str, xml_content: str, max_retries: int = 3) -> bool:
    """插入问答内容到知识库，带重试机制处理 HTTP 429。"""
    for attempt in range(max_retries):
        if attempt > 0:
            import time
            time.sleep(2 * attempt)  # 递增等待
            log.info("重试第 %d 次插入...", attempt)
        
        output = run_cli(
            "docs", "+update",
            "--doc", KB_DOC_TOKEN,
            "--command", "block_insert_after",
            "--block-id", last_block_id,
            "--content", xml_content,
            "--as", "user",
            timeout=20,
        )
        if output:
            try:
                data = json.loads(output)
                result = data.get("data", {}).get("result", "")
                if result == "success":
                    return True
                # 检查是否是频率限制
                warnings = data.get("data", {}).get("warnings", [])
                err = data.get("error", {})
                if "429" in str(err) or "429" in str(warnings) or "rate" in str(warnings).lower():
                    log.warning("遇到频率限制，等待后重试...")
                    continue
                log.warning("插入失败: %s %s", result, warnings)
                return False
            except json.JSONDecodeError:
                pass
    
    log.warning("插入失败: 已达最大重试次数 %d", max_retries)
    return False


def append_qa_to_kb(question: str, answer: str):
    """将问答对追加到知识库文档，自动累计编号。
    找到最后一个条目的答案段落，在其后插入新条目（单步插入+重试）。
    """
    with _kb_lock:
        # 实时获取文档 XML 内容
        xml_content = _fetch_kb_xml()
        if not xml_content:
            log.warning("知识库文档获取失败，跳过追加")
            return

        # 去重检查：问题前 20 个字符是否已存在
        search_text = question[:20].strip()
        if search_text and search_text in xml_content:
            log.info("问题已存在于知识库中，跳过追加: %s", search_text)
            return

        # 找到插入点（最后一个条目的答案段落 block_id）
        current_max, insert_block_id = _find_insert_point(xml_content)
        new_number = current_max + 1
        log.info("知识库当前最大编号=%d, 插入点 block=%s, 新编号=%d",
                 current_max, insert_block_id, new_number)

        if not insert_block_id:
            log.warning("未找到插入点 block ID，跳过追加")
            return

        # 构建插入内容：问题段落 + 答案段落（单步插入）
        # 答案段落首行缩进2字符（使用2个全角空格 \u3000）
        xml_content = f"<p>{new_number}\u3001{question}</p><p>\u3000\u3000{answer}</p>"

        log.info("追加问答到知识库: 编号=%d, 问题=%s", new_number, question[:40])
        success = _insert_qa_with_retry(insert_block_id, xml_content)
        if success:
            log.info("知识库追加成功: 编号=%d", new_number)
        else:
            log.warning("知识库追加失败: 编号=%d", new_number)


def extract_image_key(content: str) -> str | None:
    """从消息内容中提取图片 key（如 img_v3_xxx-xxx-xxx）。"""
    match = re.search(r'(img_[a-zA-Z0-9_-]+)', content)
    return match.group(1) if match else None


def download_image(message_id: str, image_key: str) -> str | None:
    """下载消息中的图片到临时目录，返回本地文件路径，失败返回 None。"""
    tmp_dir = tempfile.mkdtemp()
    output_path = os.path.join(tmp_dir, "image.png")
    old_cwd = os.getcwd()
    try:
        os.chdir(tmp_dir)
        output = run_cli(
            "im", "+messages-resources-download",
            "--message-id", message_id,
            "--file-key", image_key,
            "--type", "image",
            "--output", "image.png",
            "--as", "bot",
            timeout=30,
        )
    finally:
        os.chdir(old_cwd)
    if output and os.path.exists(output_path):
        log.info("图片下载成功: %s -> %s", image_key, output_path)
        return output_path
    # 尝试在 tmp_dir 中查找其他图片文件
    if os.path.exists(tmp_dir):
        for f in os.listdir(tmp_dir):
            if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
                path = os.path.join(tmp_dir, f)
                log.info("图片下载成功(其他文件名): %s -> %s", image_key, path)
                return path
    log.warning("图片下载失败: %s", image_key)
    return None


def send_image_to_user(user_id: str, image_path_or_key: str) -> str | None:
    """给指定用户发送一条纯图片消息，返回 message_id，失败返回 None。
    支持：image_key (img_xxx)、相对路径、绝对路径（自动 cd 后用相对路径）。
    """
    # 如果是绝对路径，先 cd 到目录再用相对路径（lark-cli 要求 cwd-relative）
    old_cwd = os.getcwd()
    image_arg = image_path_or_key
    try:
        if os.path.isabs(image_path_or_key) and os.path.isfile(image_path_or_key):
            img_dir = os.path.dirname(image_path_or_key)
            img_name = os.path.basename(image_path_or_key)
            os.chdir(img_dir)
            image_arg = f"./{img_name}"
            log.info("发送本地图片: cd %s && --image %s", img_dir, image_arg)

        output = run_cli(
            "im", "+messages-send",
            "--user-id", user_id,
            "--image", image_arg,
            "--as", "bot",
            "--format", "json",
        )
    finally:
        if os.getcwd() != old_cwd:
            os.chdir(old_cwd)

    if output:
        try:
            data = json.loads(output)
            mid = data.get("message_id", "") or data.get("data", {}).get("message_id", "")
            if mid:
                log.info("图片发送成功: user=%s msg_id=%s", user_id, mid)
                return mid
        except json.JSONDecodeError:
            pass
    log.warning("图片发送失败: user=%s image=%s", user_id, image_path_or_key[:50])
    return None


def upload_image(local_path: str) -> str | None:
    """上传本地图片到飞书，返回新的 image_key，失败返回 None。"""
    old_cwd = os.getcwd()
    try:
        if os.path.isabs(local_path) and os.path.isfile(local_path):
            img_dir = os.path.dirname(local_path)
            img_name = os.path.basename(local_path)
            os.chdir(img_dir)
            file_arg = f"./{img_name}"
        else:
            file_arg = local_path

        output = run_cli(
            "im", "images", "create",
            "--data", '{"image_type":"message"}',
            "--file", file_arg,
            "--as", "bot",
            "--format", "json",
            timeout=30,
        )
    finally:
        if os.getcwd() != old_cwd:
            os.chdir(old_cwd)

    if output:
        try:
            data = json.loads(output)
            key = data.get("image_key", "") or data.get("data", {}).get("image_key", "")
            if key:
                log.info("图片上传成功: %s -> %s", local_path[:40], key)
                return key
        except json.JSONDecodeError:
            pass
    log.warning("图片上传失败: %s", local_path[:50])
    return None


def send_markdown_image_to_user(user_id: str, text: str, local_image_path: str) -> str | None:
    """将文字和图片合并为一条 markdown 消息发送。
    先上传图片获取 key，再嵌入 markdown 发送。
    """
    new_key = upload_image(local_image_path)
    if not new_key:
        log.warning("图片上传失败，回退为分开发送")
        return None
    # 构建 markdown：文字 + 图片
    md_content = f"{text}\n\n![图片]({new_key})"
    mid = send_markdown_to_user(user_id, md_content)
    if mid:
        log.info("图文合并消息发送成功: user=%s", user_id)
        return mid
    log.warning("图文合并消息发送失败: user=%s", user_id)
    return None


def forward_image_to_user(user_id: str, source_message_id: str, image_key: str) -> bool:
    """下载图片并转发给指定用户。始终先下载再上传，确保 bot 有权限访问。"""
    log.info("下载图片 [%s] 后重新上传发送", image_key)
    local_path = download_image(source_message_id, image_key)
    if local_path:
        mid = send_image_to_user(user_id, local_path)
        # 清理临时文件
        try:
            os.remove(local_path)
            os.rmdir(os.path.dirname(local_path))
        except Exception:
            pass
        if mid:
            log.info("图片下载+重传成功: user=%s", user_id)
            return True
        else:
            log.warning("图片重传失败: user=%s", user_id)
    else:
        log.warning("图片下载失败，无法转发: %s", image_key)
    return False


# ─── 核心处理 ──────────────────────────────────────────────────────────────
def handle_message(event: dict):
    """处理一条收到的消息事件。"""
    message_id = event.get("message_id", "")
    chat_type = event.get("chat_type", "")
    content = event.get("content", "").strip()
    sender_id = event.get("sender_id", "")
    message_type = event.get("message_type", "text")

    if not message_id:
        log.info("skip: empty message_id")
        return
    # 图片/文件消息 content 可能为空或占位符，但仍需处理
    is_image_msg = (message_type == "image")
    if not content and not is_image_msg:
        log.info("skip: empty content (non-image)")
        return

    # 消息去重：防止同一事件被重复处理
    with _processed_lock:
        if message_id in _processed_messages:
            log.info("skip: duplicate message [%s]", message_id)
            return
        _processed_messages.add(message_id)
        # 防止集合无限增长
        if len(_processed_messages) > _MAX_PROCESSED:
            # 简单清理：保留最近一半
            to_remove = list(_processed_messages)[:_MAX_PROCESSED // 2]
            for mid in to_remove:
                _processed_messages.discard(mid)

    log.info("收到消息 [%s] from=%s chat_type=%s type=%s: %s", message_id, sender_id, chat_type, message_type, content[:100])

    # 跳过指令类消息（排除图片消息，因为图片内容以 ![Image] 开头，不是指令）
    if not is_image_msg and not content_has_image(content):
        if content.startswith("/") or content.startswith("!"):
            log.info("skip: command message")
            return

    # ── 记录最近一次提问（所有用户，非转人工关键词，始终更新） ──────────────
    # 不区分专家/普通用户：专家也可能作为提问者，需要记录其实际问题
    # 如果内容含图片引用，记录清理后的纯文本（去除 ![Image](img_xxx)）
    has_img = is_image_msg or content_has_image(content)
    clean_content = strip_image_markdown(content) if has_img else content
    if clean_content.strip() not in ("转人工", "人工", "人工客服", "转接人工"):
        _last_user_question[sender_id] = clean_content or "[图片]"
        log.info("已记录用户 [%s] 最近提问: %s", sender_id, (clean_content or "[图片]")[:50])
        # 如果是图片消息或包含嵌入图片的富文本，记录 image_key 和 message_id
        if has_img:
            img_key = extract_image_key(content)
            if img_key:
                _last_user_image[sender_id] = {"image_key": img_key, "message_id": message_id}
                log.info("已记录用户 [%s] 图片: %s (msg=%s)", sender_id, img_key, message_id)

    # ── 用户主动请求转人工 ──────────────────────────────────────────────────
    # 使用 clean_content（去除图片引用）进行关键词匹配
    if clean_content.strip() in ("转人工", "人工", "人工客服", "转接人工"):
        log.info("用户主动请求转人工")
        reply_text(message_id, "好的，已将您的需求转交人工专家处理，请稍候。")
        questioner_name = get_user_name(sender_id)
        # 使用用户最近一次的实际问题，而非"转人工"
        actual_question = _last_user_question.get(sender_id, "")
        if actual_question and actual_question not in ("[图片]",):
            expert_text = f"有用户主动请求转人工，请协助处理：\n\n提问者: {questioner_name}\n问题: {actual_question}"
        else:
            expert_text = f"有用户主动请求转人工，请协助处理：\n\n提问者: {questioner_name}\n消息: {clean_content or '[图片]'}"

        # 如果用户之前发送过图片，尝试合并为一条图文消息发送
        user_img_info = _last_user_image.get(sender_id)
        sent_msg_id = None
        if user_img_info:
            log.info("用户转人工，尝试合并图片 [%s] 发送", user_img_info["image_key"])
            local_path = download_image(user_img_info["message_id"], user_img_info["image_key"])
            if local_path:
                sent_msg_id = send_markdown_image_to_user(HUMAN_EXPERT_OPEN_ID, expert_text, local_path)
                try:
                    os.remove(local_path)
                    os.rmdir(os.path.dirname(local_path))
                except Exception:
                    pass

        # 合并发送失败或无图片，回退为纯文本
        if not sent_msg_id:
            sent_msg_id = send_text_to_user(HUMAN_EXPERT_OPEN_ID, expert_text)

        if sent_msg_id:
            with _pending_lock:
                _pending_expert_requests[sent_msg_id] = {
                    "questioner_id": sender_id,
                    "question": actual_question or clean_content or "[图片]",
                    "has_image": bool(_last_user_image.get(sender_id)),
                }
        return

    # ── 人工专家回复处理（优先级最高，含图片的回复也在此处理） ──────────────
    # 通过 +messages-mget 检查专家消息是否有 reply_to 字段
    # 有引用 → 通过 reply_to 精确匹配待处理请求，转发给对应提问者
    # 无引用 → 继续后续流程（含图片则转人工，否则知识库查询）
    if sender_id == HUMAN_EXPERT_OPEN_ID:
        reply_to = check_message_reply(message_id)
        if reply_to:
            log.info("专家引用消息回复 (reply_to=%s)", reply_to)
            with _pending_lock:
                pending = _pending_expert_requests.pop(reply_to, None)
            if pending:
                questioner_id = pending["questioner_id"]
                original_question = pending["question"]
                question_has_image = pending.get("has_image", False)
                log.info("专家回复，转发给原提问者 %s（原问题: %s）", questioner_id, original_question[:30])
                if has_img:
                    # 专家回复含图片：用 markdown 发送，自动解析 img_xxx 为图片
                    reply_msg = f"关于您的问题「{original_question[:50]}」，专家回复如下：\n\n{content}"
                    log.info("专家回复包含图片，使用 markdown 发送")
                    md_ok = send_markdown_to_user(questioner_id, reply_msg)
                    if not md_ok:
                        log.warning("markdown 发送失败，回退为纯文本")
                        send_text_to_user(questioner_id, reply_msg)
                else:
                    reply_msg = f"关于您的问题「{original_question[:50]}」，专家回复如下：\n\n{content}"
                    send_text_to_user(questioner_id, reply_msg)
                # 纯文字问题 + 专家回复：追加到知识库文档
                if not question_has_image and not has_img:
                    threading.Thread(
                        target=append_qa_to_kb,
                        args=(original_question, content),
                        daemon=True,
                    ).start()
                return
            else:
                log.info("专家引用回复但无匹配的待处理请求(reply_to=%s)，正常处理", reply_to)
        else:
            log.info("专家未引用消息，视为普通提问")

    # ── 含图片的消息：直接转人工（机器人无法处理图片内容） ──────────────
    # 所有发送者统一处理：回复提问者 + 转发文字和图片给专家（合并为一条消息）
    if has_img:
        log.info("消息包含图片，直接转人工")
        questioner_name = get_user_name(sender_id)
        question_text = clean_content if clean_content else "[图片]"
        if clean_content:
            reply_text(message_id, "收到您的图片和问题，已转交人工专家处理，请稍候。")
            expert_text = f"有用户发送图片并提问，机器人无法处理图片，请协助回复：\n\n提问者: {questioner_name}\n问题: {clean_content}"
        else:
            reply_text(message_id, "收到您的图片，已转交人工专家处理，请稍候。")
            expert_text = f"有用户发送图片，机器人无法处理，请协助回复：\n\n提问者: {questioner_name}\n问题: [图片]"

        # 下载图片，尝试与文字合并为一条 markdown 消息发送
        img_key = extract_image_key(content)
        sent_msg_id = None
        if img_key:
            log.info("下载用户图片 [%s] 用于合并发送", img_key)
            local_path = download_image(message_id, img_key)
            if local_path:
                # 合并为一条图文消息发送
                sent_msg_id = send_markdown_image_to_user(HUMAN_EXPERT_OPEN_ID, expert_text, local_path)
                # 清理临时文件
                try:
                    os.remove(local_path)
                    os.rmdir(os.path.dirname(local_path))
                except Exception:
                    pass

        # 合并发送失败，回退为分开发送
        if not sent_msg_id:
            sent_msg_id = send_text_to_user(HUMAN_EXPERT_OPEN_ID, expert_text)
            if img_key:
                log.info("回退：分开发送图片 [%s] 给专家", img_key)
                forward_image_to_user(HUMAN_EXPERT_OPEN_ID, message_id, img_key)

        if sent_msg_id:
            with _pending_lock:
                _pending_expert_requests[sent_msg_id] = {
                    "questioner_id": sender_id,
                    "question": question_text,
                    "has_image": True,
                }
        return

    # 去除图片 markdown 引用，用于知识库搜索和关键词提取
    search_text = clean_content if has_img else content

    # 1. 先抓取固定知识文档内容
    sections = []
    log.info("正在获取 %d 篇固定知识文档...", len(PINNED_DOCS))
    for doc in PINNED_DOCS:
        title = doc["title"] or "未命名文档"
        token = doc.get("token", "")
        url = doc.get("url", "")
        is_file = doc.get("is_file", False)
        local_file = doc.get("local_file", "")

        # 优先读取本地文件副本，其次用 API 下载
        if local_file:
            body = read_local_docx(local_file)
        elif is_file:
            body = fetch_file_content(token)
        else:
            body = fetch_doc_content(token)
        if body:
            summary = extract_relevant_paragraphs(body, search_text)
            if summary:
                link = f"[{title}]({url})" if url else f"**{title}**"
                sections.append(f"**{title}**\n{summary}")
                log.info("固定文档 [%s] 匹配成功 (%d 字符)", title, len(body))
            else:
                log.info("固定文档 [%s] 无强相关内容，跳过", title)
        else:
            log.warning("固定文档 [%s] 获取失败", title)

    # 2. 再搜索知识库补充更多相关文档
    log.info("搜索知识库: %s", search_text)
    search_results = search_docs(search_text)
    # 过滤掉已经在固定文档里的 token
    pinned_tokens = {d["token"] for d in PINNED_DOCS if d.get("token")}
    extra_docs = [d for d in search_results if d["token"] not in pinned_tokens]

    for doc in extra_docs[:MAX_DOCS - len(sections)] if extra_docs else []:
        title = doc["title"] or "未命名文档"
        token = doc["token"]
        url = doc.get("url", "")

        body = fetch_doc_content(token)
        if body:
            summary = extract_relevant_paragraphs(body, search_text)
            if summary:
                link = f"[{title}]({url})" if url else f"**{title}**"
                sections.append(f"**{title}**\n{summary}")
            else:
                log.info("搜索文档 [%s] 无强相关内容，跳过", title)
        else:
            log.warning("搜索文档 [%s] 获取失败，跳过", title)

    if not sections:
        # ── 转人工处理 ──────────────────────────────────────────────────────
        log.info("未找到相关文档，转人工处理（专家: %s）", HUMAN_EXPERT_OPEN_ID)
        # 通知提问者
        reply_text(message_id, "抱歉，未在知识库中找到相关答案。已将您的问题转交人工专家处理，请稍候。")
        # 转发问题给专家：优先使用 _last_user_question 中记录的实际问题
        questioner_name = get_user_name(sender_id)
        forwarded_question = _last_user_question.get(sender_id, search_text or "[图片]")
        expert_text = f"有用户提问，机器人无法回答，请协助回复：\n\n提问者: {questioner_name}\n问题: {forwarded_question}"

        # 如果用户发送了图片，尝试合并为一条图文消息发送
        user_img_info = _last_user_image.get(sender_id)
        sent_msg_id = None
        if user_img_info:
            log.info("用户自动转人工，尝试合并图片 [%s] 发送", user_img_info["image_key"])
            local_path = download_image(user_img_info["message_id"], user_img_info["image_key"])
            if local_path:
                sent_msg_id = send_markdown_image_to_user(HUMAN_EXPERT_OPEN_ID, expert_text, local_path)
                try:
                    os.remove(local_path)
                    os.rmdir(os.path.dirname(local_path))
                except Exception:
                    pass

        # 合并发送失败或无图片，回退为纯文本
        if not sent_msg_id:
            sent_msg_id = send_text_to_user(HUMAN_EXPERT_OPEN_ID, expert_text)

        if sent_msg_id:
            with _pending_lock:
                _pending_expert_requests[sent_msg_id] = {
                    "questioner_id": sender_id,
                    "question": forwarded_question,
                    "has_image": False,
                }
        return

    # 3. 组装回复
    # 使用 search_text（已去除图片引用）展示用户问题
    display_text = search_text if search_text else "[图片]"
    header = f"根据你的问题「{display_text[:50]}」，找到以下相关内容：\n\n"
    reply = header + "\n\n".join(sections) + TRANSFER_FOOTER

    log.info("发送回复 (%d 字符)", len(reply))
    ok = reply_message(message_id, reply)
    if ok:
        log.info("回复成功")
    else:
        # 降级为纯文本
        log.warning("markdown 回复失败，降级为纯文本")
        plain = header.replace("**", "").replace("[", "").replace("](", " (").replace(")", "")
        for s in sections:
            plain += s.replace("### ", "").replace("**", "") + "\n---\n"
        plain += "\n若回复未能解答您的疑问，请直接回复「转人工」，我们将为您转交专家处理。"
        reply_text(message_id, plain)


# ─── 主循环 ────────────────────────────────────────────────────────────────
def main():
    log.info("启动知识库问答机器人...")
    log.info("lark-cli: %s", LARK_CLI)
    log.info("知识库空间: %s", WIKI_SPACE_IDS or "全部")

    # 启动 event consume，通过 stdin 保持连接（写入空行保持 stdin 不关闭）
    consume_cmd = [
        LARK_CLI,
        "event", "consume", "im.message.receive_v1",
        "--as", "bot",
    ]

    log.info("开始监听消息事件: %s", " ".join(consume_cmd))

    # 用 Popen 持续读取 stdout
    proc = subprocess.Popen(
        consume_cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        stdin=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        bufsize=1,
    )

    # 启动 stderr 读取线程（打印日志）
    def read_stderr():
        for line in proc.stderr:
            line = line.strip()
            if line:
                log.info("[event-consume] %s", line)
                if "ready" in line:
                    log.info("事件监听已就绪，等待用户消息...")

    stderr_thread = threading.Thread(target=read_stderr, daemon=True)
    stderr_thread.start()

    # 保持 stdin 打开（写入空内容防止 EOF）
    try:
        proc.stdin.write("\n")
        proc.stdin.flush()
    except Exception:
        pass

    # 主循环：逐行读取事件
    try:
        for line in proc.stdout:
            line = line.strip()
            if not line:
                continue
            try:
                event = json.loads(line)
                # 用线程处理，避免阻塞事件消费
                threading.Thread(
                    target=handle_message, args=(event,), daemon=True
                ).start()
            except json.JSONDecodeError:
                log.warning("非解析事件: %s", line[:200])
    except KeyboardInterrupt:
        log.info("收到中断信号，正在停止...")
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()
        log.info("机器人已停止")


if __name__ == "__main__":
    main()
